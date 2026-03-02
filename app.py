import io
import textwrap
from datetime import datetime

import streamlit as st
from docx import Document
from openai import OpenAI
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas


API_KEY = "sk-99a896646dc0498fa779c88df07530c0"
BASE_URL = "https://dashscope.aliyuncs.com/compatible-mode/v1"
MODEL_NAME = "qwen-plus"

SYSTEM_PROMPT = (
    "你是一个资深大厂HR。请对比用户的简历和岗位的JD，犀利指出简历中欠缺的关键词，"
    "并重写一段更亮眼、更匹配该岗位的个人优势或经历描述。"
)

DOWNLOAD_FORMAT_OPTIONS = (
    "纯文本文件（.txt）",
    "文档文件（.docx）",
    "便携文档（.pdf）",
)


def extract_uploaded_jd(uploaded_file) -> str:
    extension = uploaded_file.name.rsplit(".", 1)[-1].lower()

    if extension == "txt":
        raw = uploaded_file.getvalue()
        for encoding in ("utf-8", "utf-8-sig", "gbk", "gb18030"):
            try:
                return raw.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
        raise ValueError("TXT 文件编码无法识别，请使用 UTF-8 或 GBK 编码。")

    if extension == "docx":
        uploaded_file.seek(0)
        document = Document(uploaded_file)
        text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
        if not text.strip():
            raise ValueError("DOCX 文件中未识别到可用文字内容。")
        return text.strip()

    if extension == "pdf":
        uploaded_file.seek(0)
        reader = PdfReader(uploaded_file)
        text = "\n".join((page.extract_text() or "") for page in reader.pages).strip()
        if not text:
            raise ValueError("PDF 文件未提取到文字，可能是扫描图片版 PDF。")
        return text

    raise ValueError("仅支持上传 .txt、.docx、.pdf 文件。")


def call_llm(jd_text: str, resume_text: str) -> str:
    client = OpenAI(api_key=API_KEY, base_url=BASE_URL)
    user_prompt = f"""
【目标岗位的职位描述（JD）】
{jd_text}

【求职者目前的简历经历】
{resume_text}

请按以下结构输出：
1) 简历中欠缺的关键词（按重要性排序，尽量具体）
2) 重写后的亮眼个人优势或经历描述（可直接粘贴进简历）
3) 3 条可立即执行的改写建议（含“如何量化成果”的写法）
""".strip()

    response = client.chat.completions.create(
        model=MODEL_NAME,
        temperature=0.5,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt},
        ],
    )
    content = response.choices[0].message.content
    if not content:
        raise ValueError("模型返回为空，请稍后重试。")
    return content.strip()


def build_txt_bytes(content: str) -> bytes:
    return content.encode("utf-8-sig")


def build_docx_bytes(content: str) -> bytes:
    buffer = io.BytesIO()
    document = Document()
    document.add_heading("简历优化结果", level=1)
    for line in content.splitlines():
        document.add_paragraph(line if line.strip() else "")
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def build_pdf_bytes(content: str) -> bytes:
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    page_width, page_height = A4

    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        font_name = "STSong-Light"
    except Exception:
        font_name = "Helvetica"

    pdf.setFont(font_name, 14)
    y = page_height - 52
    pdf.drawString(48, y, "简历优化结果")
    y -= 28
    pdf.setFont(font_name, 11)

    for paragraph in content.splitlines():
        wrapped_lines = textwrap.wrap(
            paragraph,
            width=36,
            break_long_words=True,
            drop_whitespace=False,
        ) or [""]
        for line in wrapped_lines:
            if y < 52:
                pdf.showPage()
                pdf.setFont(font_name, 11)
                y = page_height - 52
            pdf.drawString(48, y, line)
            y -= 18

    pdf.save()
    buffer.seek(0)
    return buffer.getvalue()


def build_download_payload(content: str, download_type: str) -> tuple[bytes, str, str]:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    if download_type == "纯文本文件（.txt）":
        return build_txt_bytes(content), "text/plain", f"简历优化结果_{timestamp}.txt"
    if download_type == "文档文件（.docx）":
        return (
            build_docx_bytes(content),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            f"简历优化结果_{timestamp}.docx",
        )
    return build_pdf_bytes(content), "application/pdf", f"简历优化结果_{timestamp}.pdf"


st.set_page_config(
    page_title="🚀 大学生秋招简历与岗位匹配神器",
    page_icon="🚀",
    layout="wide",
)

st.markdown(
    """
<style>
    :root {
        --bg-main: #f6fbff;
        --bg-card: rgba(255, 255, 255, 0.84);
        --text-title: #152c4a;
        --text-normal: #38506d;
        --accent-a: #1e80ff;
        --accent-b: #ff8a3d;
        --line-soft: rgba(22, 62, 104, 0.14);
    }
    html, body, [class*="css"] {
        font-family: "MiSans", "HarmonyOS Sans SC", "Source Han Sans SC", "PingFang SC", "Microsoft YaHei", sans-serif;
    }
    .stApp {
        background:
            radial-gradient(circle at 8% 12%, rgba(30, 128, 255, 0.22), transparent 36%),
            radial-gradient(circle at 88% 24%, rgba(255, 138, 61, 0.18), transparent 34%),
            radial-gradient(circle at 52% 90%, rgba(76, 216, 255, 0.16), transparent 36%),
            var(--bg-main);
    }
    .block-container {
        padding-top: 1.2rem;
        max-width: 1080px;
    }
    .hero {
        border-radius: 20px;
        padding: 1.3rem 1.5rem;
        background: linear-gradient(118deg, rgba(255,255,255,0.95) 0%, rgba(245,251,255,0.9) 60%, rgba(255,247,239,0.92) 100%);
        border: 1px solid var(--line-soft);
        box-shadow: 0 16px 38px rgba(29, 61, 97, 0.12);
        backdrop-filter: blur(10px);
        margin-bottom: 1.2rem;
        animation: slideIn 0.45s ease-out;
    }
    .hero h1 {
        margin: 0;
        color: var(--text-title);
        font-size: 1.9rem;
        line-height: 1.25;
    }
    .hero p {
        margin: 0.55rem 0 0 0;
        color: var(--text-normal);
        font-size: 1.02rem;
        letter-spacing: 0.01em;
    }
    .chip-row {
        display: flex;
        flex-wrap: wrap;
        gap: 0.55rem;
        margin-top: 0.9rem;
    }
    .chip {
        border-radius: 999px;
        padding: 0.34rem 0.8rem;
        font-size: 0.85rem;
        color: #21446e;
        border: 1px solid rgba(30, 128, 255, 0.2);
        background: rgba(232, 243, 255, 0.72);
    }
    .result-card {
        margin-top: 0.75rem;
        padding: 1.15rem 1.2rem;
        border-radius: 16px;
        background: #ffffff;
        border: 1px solid rgba(46, 111, 219, 0.16);
        box-shadow: 0 12px 28px rgba(23, 43, 77, 0.09);
    }
    .section-title {
        margin: 0.2rem 0 0.65rem 0;
        color: var(--text-title);
        font-size: 1.05rem;
        font-weight: 700;
    }
    .section-sub {
        margin: 0 0 0.85rem 0;
        color: #5c7391;
        font-size: 0.9rem;
    }
    div[data-testid="stTextArea"] label p {
        color: #1e3d63;
        font-weight: 700;
        font-size: 1rem;
    }
    div[data-testid="stTextArea"] textarea {
        background: rgba(255, 255, 255, 0.92);
        border-radius: 14px;
        border: 1px solid rgba(46, 111, 219, 0.22);
        box-shadow: inset 0 1px 0 rgba(255,255,255,0.7);
        color: #1f3552;
        transition: border-color .2s ease, box-shadow .2s ease, transform .2s ease;
    }
    div[data-testid="stTextArea"] textarea:focus {
        border-color: rgba(30, 128, 255, 0.7);
        box-shadow: 0 0 0 0.2rem rgba(30, 128, 255, 0.13);
        transform: translateY(-1px);
    }
    div[data-testid="stFileUploaderDropzone"] {
        background: rgba(255, 255, 255, 0.92);
        border: 1.5px dashed rgba(46, 111, 219, 0.4);
        border-radius: 16px;
        padding: 0.8rem 0.8rem 1rem 0.8rem;
    }
    div[data-testid="stFileUploaderDropzoneInstructions"] > div {
        display: none;
    }
    div[data-testid="stFileUploaderDropzoneInstructions"]::before {
        content: "将职位描述文件拖拽到此处";
        display: block;
        color: #1f4874;
        font-weight: 600;
        margin-bottom: 0.15rem;
    }
    div[data-testid="stFileUploaderDropzone"] small {
        display: none;
    }
    div[data-testid="stFileUploaderDropzone"] button {
        font-size: 0 !important;
        border-radius: 999px;
        border: none !important;
        color: transparent !important;
        background: linear-gradient(92deg, #1e80ff 0%, #4f9cff 100%) !important;
        box-shadow: 0 8px 18px rgba(30, 128, 255, 0.25);
    }
    div[data-testid="stFileUploaderDropzone"] button::after {
        content: "选择本地文件";
        color: white;
        font-size: 0.9rem;
        font-weight: 700;
    }
    div[data-testid="stFileUploaderDropzone"]::after {
        content: "支持格式：txt、docx、pdf";
        display: block;
        margin-top: 0.45rem;
        color: #59718f;
        font-size: 0.84rem;
    }
    div[data-testid="stButton"] > button {
        border: none;
        border-radius: 14px;
        background: linear-gradient(95deg, var(--accent-a) 0%, #4a8cff 52%, var(--accent-b) 100%);
        color: #ffffff;
        font-weight: 800;
        letter-spacing: 0.02em;
        padding-top: 0.65rem;
        padding-bottom: 0.65rem;
        box-shadow: 0 12px 26px rgba(30, 128, 255, 0.32);
        transition: transform .2s ease, box-shadow .2s ease;
    }
    div[data-testid="stButton"] > button:hover {
        transform: translateY(-1px);
        box-shadow: 0 16px 28px rgba(30, 128, 255, 0.35);
    }
    div[data-baseweb="select"] > div {
        border-radius: 12px;
        border: 1px solid rgba(46, 111, 219, 0.25);
        background: rgba(255, 255, 255, 0.9);
    }
    @keyframes slideIn {
        from { opacity: 0; transform: translateY(8px); }
        to { opacity: 1; transform: translateY(0); }
    }
    @media (max-width: 768px) {
        .hero h1 {
            font-size: 1.55rem;
        }
        .hero p {
            font-size: 0.93rem;
        }
        .block-container {
            padding-top: 0.8rem;
            padding-left: 1rem;
            padding-right: 1rem;
        }
    }
</style>
""",
    unsafe_allow_html=True,
)

st.markdown(
    """
<div class="hero">
  <h1>🚀 大学生秋招简历与岗位匹配神器</h1>
  <p>一键对齐心仪岗位要求，资深人力资源视角帮你精准润色简历经历。</p>
  <div class="chip-row">
    <span class="chip">关键词差距诊断</span>
    <span class="chip">亮点经历重写</span>
    <span class="chip">学生友好中文界面</span>
  </div>
</div>
""",
    unsafe_allow_html=True,
)

if "jd_text" not in st.session_state:
    st.session_state.jd_text = ""
if "resume_text" not in st.session_state:
    st.session_state.resume_text = ""
if "latest_result" not in st.session_state:
    st.session_state.latest_result = ""
if "last_uploaded_file_token" not in st.session_state:
    st.session_state.last_uploaded_file_token = ""

uploaded_jd_file = st.file_uploader(
    "上传岗位说明文件（可选，上传后会自动填充左侧输入框）",
    type=["txt", "docx", "pdf"],
)
if uploaded_jd_file is not None:
    current_token = f"{uploaded_jd_file.name}-{uploaded_jd_file.size}"
    if current_token != st.session_state.last_uploaded_file_token:
        try:
            st.session_state.jd_text = extract_uploaded_jd(uploaded_jd_file)
            st.session_state.last_uploaded_file_token = current_token
            st.success("已成功读取上传文件，岗位描述内容已自动填充。")
        except Exception as exc:
            st.error(f"文件读取失败：{exc}")

st.markdown('<p class="section-title">请填写岗位要求与简历内容</p>', unsafe_allow_html=True)
st.markdown('<p class="section-sub">左侧输入岗位说明，右侧输入你的当前简历；系统会自动给出更贴合岗位的改写建议。</p>', unsafe_allow_html=True)

col_left, col_right = st.columns(2, gap="large")
with col_left:
    st.text_area(
        "目标岗位职位描述",
        key="jd_text",
        height=320,
        placeholder="示例：负责用户增长策略分析、活动运营与数据复盘，要求掌握结构化查询、数据分析、分组实验方法...",
    )
with col_right:
    st.text_area(
        "目前的简历经历",
        key="resume_text",
        height=320,
        placeholder="示例：教育背景、实习经历、项目经历、专业技能与可量化成果...",
    )

if st.button("✨ 一键优化简历", type="primary", use_container_width=True):
    jd_text = st.session_state.jd_text.strip()
    resume_text = st.session_state.resume_text.strip()
    if not jd_text or not resume_text:
        st.warning("请先完整填写左侧岗位描述和右侧简历经历，再进行优化。")
    else:
        with st.spinner("资深招聘顾问正在分析中，请稍候..."):
            try:
                st.session_state.latest_result = call_llm(jd_text, resume_text)
            except Exception as exc:
                st.error(f"请求失败：{exc}")

if st.session_state.latest_result:
    st.markdown("### 智能优化结果")
    st.markdown('<div class="result-card">', unsafe_allow_html=True)
    st.markdown(st.session_state.latest_result)
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("#### 下载优化结果")
    download_type = st.selectbox(
        "请选择下载格式",
        DOWNLOAD_FORMAT_OPTIONS,
        index=0,
    )
    file_bytes, mime, file_name = build_download_payload(st.session_state.latest_result, download_type)
    st.download_button(
        "下载当前优化结果",
        data=file_bytes,
        file_name=file_name,
        mime=mime,
        use_container_width=True,
    )
